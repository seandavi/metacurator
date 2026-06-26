"""tables tests (SPEC 050) — offline, fixtures generated at runtime."""

from __future__ import annotations

import pytest

from metacurator.models import SourceProvenance
from metacurator.tables import _clean_headers, _rows_to_table, load_tables

# -- shared row->table conversion (also the pdf code path) -------------------


def test_rows_to_table_header_detection():
    raw = [
        ["Supplementary Table 1", None],          # title row, mostly empty
        ["subject_id", "disease"],                # real header (row 1)
        ["s1", "Healthy"],
        ["s2", "Crohn Disease"],
    ]
    table = _rows_to_table(raw, SourceProvenance(file="x"))
    assert table.provenance.header_row == 1
    assert table.frame.columns == ["subject_id", "disease"]
    assert table.frame.records[0] == {"subject_id": "s1", "disease": "Healthy"}
    assert table.n_rows == 2 and table.n_cols == 2


def test_clean_blank_and_duplicate_headers():
    assert _clean_headers(["id", "", "id"]) == ["id", "col2", "id_1"]
    assert _clean_headers(["a", None, "a", "a"]) == ["a", "col2", "a_1", "a_2"]


def test_rows_to_table_empty_returns_none():
    assert _rows_to_table([[None, None], ["", ""]], SourceProvenance(file="x")) is None


# -- delimited ---------------------------------------------------------------


def test_csv_and_tsv_roundtrip(tmp_path):
    rows = "subject_id,disease\ns1,Healthy\ns2,Crohn Disease\n"
    (tmp_path / "a.csv").write_text(rows)
    (tmp_path / "a.tsv").write_text(rows.replace(",", "\t"))
    csv_tables = load_tables(tmp_path / "a.csv")
    tsv_tables = load_tables(tmp_path / "a.tsv")
    assert csv_tables[0].frame.records == tsv_tables[0].frame.records
    assert csv_tables[0].frame.records[1] == {"subject_id": "s2", "disease": "Crohn Disease"}


def test_no_tables_is_error(tmp_path):
    (tmp_path / "empty.csv").write_text("\n\n")
    with pytest.raises(ValueError, match="no tables"):
        load_tables(tmp_path / "empty.csv")


def test_unsupported_suffix(tmp_path):
    p = tmp_path / "x.parquet"
    p.write_text("nope")
    with pytest.raises(ValueError, match="unsupported"):
        load_tables(p)


# -- xlsx --------------------------------------------------------------------


def test_xlsx_multisheet_and_header_detection(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    s1 = wb.active
    s1.title = "samples"
    s1.append(["subject_id", "disease"])
    s1.append(["s1", "Healthy"])
    s2 = wb.create_sheet("with_title")
    s2.append(["A big title", None])      # header not on row 0
    s2.append(["country", "bmi"])
    s2.append(["Italy", "23.5"])
    path = tmp_path / "supp.xlsx"
    wb.save(path)

    tables = load_tables(path)
    by_sheet = {t.provenance.sheet: t for t in tables}
    assert set(by_sheet) == {"samples", "with_title"}
    assert by_sheet["samples"].provenance.header_row == 0
    assert by_sheet["samples"].frame.records[0]["disease"] == "Healthy"
    assert by_sheet["with_title"].provenance.header_row == 1
    assert by_sheet["with_title"].frame.columns == ["country", "bmi"]


# -- docx --------------------------------------------------------------------


def test_docx_table(tmp_path):
    from docx import Document

    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    for col, name in enumerate(["subject_id", "disease"]):
        t.cell(0, col).text = name
    t.cell(1, 0).text, t.cell(1, 1).text = "s1", "Healthy"
    t.cell(2, 0).text, t.cell(2, 1).text = "s2", "Crohn Disease"
    path = tmp_path / "supp.docx"
    doc.save(path)

    tables = load_tables(path)
    assert len(tables) == 1
    assert tables[0].provenance.table_index == 0
    assert tables[0].frame.records[1] == {"subject_id": "s2", "disease": "Crohn Disease"}


# -- pdf (live parsing gated on a PDF writer being available) ----------------


def test_pdf_table(tmp_path):
    pytest.importorskip("reportlab", reason="no PDF writer to build a fixture")
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table

    path = tmp_path / "supp.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    data = [["subject_id", "disease"], ["s1", "Healthy"], ["s2", "Crohn Disease"]]
    doc.build([Table(data)])

    tables = load_tables(path)
    assert tables
    recs = tables[0].frame.records
    assert {"s1", "s2"} <= {r.get("subject_id") for r in recs}
